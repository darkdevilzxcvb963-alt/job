"""
Resume Parsing Service
Handles extraction of text and structured data from PDF and DOCX files
"""
import os
from typing import Dict, Optional
from pathlib import Path
import pdfplumber
from docx import Document
import time
import requests
from loguru import logger
from app.core.config import settings

# Optional imports for fallback extractors

# Optional imports for fallback extractors
try:
    import fitz as pymupdf
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import pypdf
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    import pytesseract
    from pdf2image import convert_from_path
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

try:
    import pypdfium2
    HAS_PYPDFIUM = True
except ImportError:
    HAS_PYPDFIUM = False

class ResumeParser:
    """Parse resumes from various file formats"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc']
    
    def parse(self, file_path: str) -> Dict[str, any]:
        """
        Parse resume file and extract text
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Dictionary with parsed data including text and metadata
        """
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return self._parse_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _parse_pdf(self, file_path: str) -> Dict[str, any]:
        """Parse PDF file using multiple extraction methods as fallback"""
        full_text = ""
        method_used = "none"
        
        # Method 1: pdfplumber (best for text-based PDFs)
        try:
            text_content = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
            full_text = "\n".join(text_content)
            if full_text.strip():
                method_used = "pdfplumber"
        except Exception as e:
            logger.warning(f"pdfplumber failed for {file_path}: {e}")
        
        # Method 2: PyMuPDF fallback
        if not full_text.strip() and HAS_PYMUPDF:
            try:
                # fitz / pymupdf
                doc = pymupdf.open(file_path)
                text_parts = []
                for page in doc:
                    text = page.get_text()
                    if text:
                        text_parts.append(text)
                doc.close()
                full_text = "\n".join(text_parts)
                if full_text.strip():
                    method_used = "pymupdf"
            except Exception as e:
                logger.warning(f"PyMuPDF failed for {file_path}: {e}")
        
        # Method 3: pypdf (modern replacement for PyPDF2)
        if not full_text.strip() and HAS_PYPDF:
            try:
                reader = pypdf.PdfReader(file_path)
                text_parts = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                full_text = "\n".join(text_parts)
                if full_text.strip():
                    method_used = "pypdf"
            except Exception as e:
                logger.warning(f"pypdf failed for {file_path}: {e}")
        
        # Method 4: pypdfium2 fallback (very fast and robust)
        if not full_text.strip() and HAS_PYPDFIUM:
            try:
                pdf = pypdfium2.PdfDocument(file_path)
                text_parts = []
                for i in range(len(pdf)):
                    page = pdf[i]
                    textpage = page.get_textpage()
                    text = textpage.get_text_range()
                    if text:
                        text_parts.append(text)
                full_text = "\n".join(text_parts)
                if full_text.strip():
                    method_used = "pypdfium2"
            except Exception as e:
                logger.warning(f"pypdfium2 failed for {file_path}: {e}")

        # Method 5: OCR fallback (for scanned/image PDFs)
        if not full_text.strip() and HAS_OCR:
            try:
                logger.info(f"Attempting OCR for scanned PDF: {file_path}")
                images = convert_from_path(file_path)
                text_parts = []
                for img in images:
                    text = pytesseract.image_to_string(img)
                    if text:
                        text_parts.append(text)
                full_text = "\n".join(text_parts)
                if full_text.strip():
                    method_used = "ocr"
            except Exception as e:
                logger.warning(f"OCR failed for {file_path}: {e}")
        
        # Method 6: LlamaParse fallback (Elite AI parsing)
        if not full_text.strip() and settings.LLAMA_CLOUD_API_KEY:
            try:
                logger.info(f"Attempting LlamaParse for complex/scanned PDF: {file_path}")
                llamaparse_text = self._parse_with_llamaparse(file_path)
                if llamaparse_text:
                    full_text = llamaparse_text
                    method_used = "llamaparse"
            except Exception as e:
                logger.warning(f"LlamaParse failed for {file_path}: {e}")
        
        extraction_warning = None
        if not full_text.strip():
            if not settings.LLAMA_CLOUD_API_KEY:
                extraction_warning = "No text found. This appears to be a scanned or complex PDF. Please provide a LLAMA_CLOUD_API_KEY for advanced AI parsing."
            else:
                extraction_warning = "All extraction methods, including LlamaParse, returned empty text."
        
        if full_text.strip():
            logger.info(f"PDF parsed successfully with {method_used}: {len(full_text)} chars from {file_path}")
        else:
            logger.warning(f"All PDF extraction methods returned empty text for {file_path}. {extraction_warning}")
        
        return {
            "text": full_text,
            "file_type": "pdf",
            "extraction_method": method_used,
            "warning": extraction_warning
        }
    
    def _parse_with_llamaparse(self, file_path: str) -> Optional[str]:
        """Parse PDF using LlamaParse Direct REST API"""
        if not settings.LLAMA_CLOUD_API_KEY:
            return None
            
        try:
            url = "https://api.cloud.llamaindex.ai/api/parsing/upload"
            headers = {"Authorization": f"Bearer {settings.LLAMA_CLOUD_API_KEY}"}
            
            with open(file_path, "rb") as f:
                files = {"file": f}
                data = {"result_type": "markdown"}
                response = requests.post(url, headers=headers, files=files, data=data, timeout=60)
                
            if response.status_code != 200:
                logger.error(f"LlamaParse upload failed: {response.text}")
                return None
                
            job_id = response.json().get("id")
            if not job_id:
                return None
                
            # Polling for results
            max_retries = 60 # 60 seconds max
            for _ in range(max_retries):
                status_url = f"https://api.cloud.llamaindex.ai/api/parsing/job/{job_id}"
                status_res = requests.get(status_url, headers=headers, timeout=10)
                status_data = status_res.json()
                
                if status_data.get("status") == "SUCCESS":
                    result_url = f"https://api.cloud.llamaindex.ai/api/parsing/job/{job_id}/result/markdown"
                    result_res = requests.get(result_url, headers=headers, timeout=30)
                    return result_res.json().get("markdown")
                elif status_data.get("status") == "ERROR":
                    logger.error(f"LlamaParse job error: {status_data}")
                    return None
                    
                time.sleep(1)
                
            logger.warning("LlamaParse timeout reached")
            return None
        except Exception as e:
            logger.error(f"LlamaParse REST API error: {e}")
            return None
    
    def _parse_docx(self, file_path: str) -> Dict[str, any]:
        """Parse DOCX file including text from tables, paragraphs, headers, and footers"""
        try:
            doc = Document(file_path)
            text_content = []
            
            # 1. Extract from headers
            for section in doc.sections:
                header = section.header
                for p in header.paragraphs:
                    if p.text.strip():
                        text_content.append(p.text.strip())
            
            # 2. Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # 3. Extract from tables (resumes often use tables for layout)
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        # Extract nested paragraphs in cells
                        cell_parts = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
                        cell_text = " ".join(cell_parts)
                        if cell_text and cell_text not in row_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            # 4. Extract from footers
            for section in doc.sections:
                footer = section.footer
                for p in footer.paragraphs:
                    if p.text.strip():
                        text_content.append(p.text.strip())
            
            full_text = "\n".join(text_content)
            
            # Fallback if text is still very short
            if len(full_text.strip()) < 100:
                 # Check for text boxes or other elements? (Limited support in python-docx)
                 pass

            return {
                "text": full_text,
                "file_type": "docx",
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections)
            }
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            # If it's an old .doc file, python-docx will raise a PackageNotFoundError or similar
            if "not a Word file" in str(e) or "zip" in str(e).lower():
                 return {
                    "text": "",
                    "file_type": "docx-incompatible",
                    "error": "Old .doc files are not supported. Please convert to .pdf or .docx"
                 }
            raise


_resume_parser = None

def get_resume_parser():
    """Factory function for ResumeParser"""
    global _resume_parser
    if _resume_parser is None:
        _resume_parser = ResumeParser()
    return _resume_parser

